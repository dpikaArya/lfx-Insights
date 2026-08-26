"""Render a grounded :class:`SectionBundle` to an APA-styled .docx.

``python-docx`` is imported lazily inside the functions so importing this module
(and the CLI) never requires the optional ``[docx]`` extra; only calling the
render path does.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

from lfx_insights.errors import ConsiliumError
from lfx_insights.generation.common import disambiguation_suffixes, format_apa

if TYPE_CHECKING:
    from lfx_insights.models import SectionBundle

_DOCX_HINT = "DOCX export needs the 'docx' extra: pip install consilium[docx]"


def _document() -> Any:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - exercised only without the extra
        raise ConsiliumError(_DOCX_HINT) from exc
    return Document()


def render_docx(bundle: SectionBundle) -> Any:
    """Build and return a python-docx ``Document`` for ``bundle``."""
    doc = _document()
    doc.add_heading(bundle.title, level=1)

    for section in bundle.sections:
        doc.add_heading(section.name.replace("_", " ").title(), level=1)
        for para in section.text.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    doc.add_heading("References (APA)", level=1)
    if not bundle.references:
        doc.add_paragraph("No references.")
        return doc

    # Same disambiguation map as the in-text citations, so a "Smith et al., 2020a"
    # in the prose matches its "2020a" reference entry.
    suffixes = disambiguation_suffixes(bundle.references)
    for ref in sorted(format_apa(p, suffixes.get(p.id, "")) for p in bundle.references):
        paragraph = doc.add_paragraph(ref)
        paragraph.paragraph_format.left_indent = _hanging_indent()
        paragraph.paragraph_format.first_line_indent = -_hanging_indent()
    return doc


def _hanging_indent() -> Any:
    from docx.shared import Inches

    return Inches(0.5)


def write_docx(bundle: SectionBundle, path: str | Path) -> Path:
    """Render ``bundle`` and save it to ``path``; returns the path."""
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    render_docx(bundle).save(str(out))
    return out
